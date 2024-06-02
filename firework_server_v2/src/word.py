import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx.shared import Pt
from io import BytesIO

def word_create(card, client):

    doc = docx.Document()

    c = 0
    doc.add_heading('Заключение', 0)
    doc.paragraphs[c].runs[0].bold = True
    doc.paragraphs[c].style = "Heading 1"
    doc.paragraphs[c].runs[0].font.size = Pt(16)
    doc.paragraphs[c].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Пациент ')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.paragraphs[c].add_run(card.clientFullName)

    doc.add_paragraph('обратился ')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.paragraphs[c].add_run(client.created.strftime('%Y-%m-%d'))

    # Жалобы пациента
    doc.add_heading('Общая картина на основе консультации:', 2)
    c += 1
    doc.paragraphs[c].runs[0].bold = True
    doc.paragraphs[c].style = "Heading 2"
    doc.paragraphs[c].runs[0].font.size = Pt(14)


    doc.add_paragraph('Психоэмоциональное состояние:')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.comPsychoEmotionNote)
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Соматические жалобы:')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.comSomatic)
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Социальные жалобы:')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.comSocial)
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    # Анамнез заболевания
    doc.add_heading('В анамнезе отмечается:', 2)
    c += 1
    doc.paragraphs[c].runs[0].bold = True
    doc.paragraphs[c].style = "Heading 2"
    doc.paragraphs[c].runs[0].font.size = Pt(14)

    doc.add_paragraph('Пациент живет с онкологическим заболеванием:')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.oncologyTime)
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Эмоциональные реакции в момент первого озвучивания диагноза:')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.emotionReactOnco)
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Справлялся с реакциями:')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.experiencingEmotion)
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Пациенту обычно помогает справляться с эмоциональными реакциями:')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.usuallyHelpEmotion)
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Не помогает справляться :')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.usuallyNoHelpEmotion)
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Связь с психотравмой: ')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasConnectionPsychotrauma:
        doc.add_paragraph("Есть.")
    else:
        doc.add_paragraph("Нет.")
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.paragraphs[c].add_run(card.connectionPsychotraumaNote if card.connectionPsychotraumaNote else ' ')
    doc.paragraphs[c].runs[1].font.size = Pt(12)

    # Раздел “Объективный статус”
    doc.add_heading('Объективно:', 2)
    c += 1
    doc.paragraphs[c].runs[0].bold = True
    doc.paragraphs[c].style = "Heading 2"
    doc.paragraphs[c].runs[0].font.size = Pt(14)

    doc.add_paragraph('В беседе: ')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.paragraphs[c].add_run(card.featuresConversationsNote if card.featuresConversationsNote else ' ')
    doc.paragraphs[c].runs[1].font.size = Pt(12)

    doc.add_paragraph('Признаки тревоги: ')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph('Есть.' if card.hasAnxiety else 'Нет.')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Депрессивные симптомы: ')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasContinuousDuration:
        doc.add_paragraph('\t- Тервога более 2 недель')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasPersistentLowMood:
        doc.add_paragraph('\t- Постоянное снижение настроения')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasLossInterest:
        doc.add_paragraph('\t- Утрата интересов')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasEnergyDecline:
        doc.add_paragraph('\t- Снижение энергии или усталось после минимальных усилий')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasDecliningAttention:
        doc.add_paragraph('\t- Снижение концентрации внимания')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasGuilt:
        doc.add_paragraph('\t- Чувство вины и никчемность')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasLowerSelfesteem:
        doc.add_paragraph('\t- Снижение самооценки и уверенности в себе')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasAppetiteChange:
        doc.add_paragraph('\t- Изменение аппетита и веса')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasSleepDisturbance:
        doc.add_paragraph('\t- Нарушение сна')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasPessimisticThoughts:
        doc.add_paragraph('\t- Пессимистические мысли и погружение в себя')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasHopelessness:
        doc.add_paragraph('\t- Безнадежность')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasSuicidalThoughts:
        doc.add_paragraph('\t- Суицидальные мысли и действия')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Влияние симптомов на межличностные отношения: ')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.symptomAtRelationshipNote if card.symptomAtRelationshipNote else 'Отсутствует.')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Суицидальные мысли: ')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.suicidalThoughtsNote if card.suicidalThoughtsNote else 'Отсутствуют.')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)


    # Раздел “Внутренняя картина болезни”
    doc.add_heading('Внутренняя картина болезни характеризуется следующим образом:', 2)
    c += 1
    doc.paragraphs[c].runs[0].bold = True
    doc.paragraphs[c].style = "Heading 2"
    doc.paragraphs[c].runs[0].font.size = Pt(14)

    doc.add_paragraph('Отношение пациента к болезни:')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.opinionPersonIllnessNote if card.opinionPersonIllnessNote else 'Нет.')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Пациент связывает начало болезни с:')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.opinionOnsetOfDiseaseNote if card.opinionOnsetOfDiseaseNote else 'Нет.')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Пациент относится к лечению:')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.opinionTreatmentNote if card.opinionTreatmentNote else 'Нет.')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Пациент думает о будущем:')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.opinionFutureNote if card.opinionFutureNote else 'Нет.')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Копинги ведущие:')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.copingHostsNote if card.copingHostsNote else 'Нет.')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Активные копинги: ')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph('Наблюдаются' if card.hasActiveCopingSpecies else 'Не наблюдаются')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Копинги развиты: ')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph('Хорошо' if card.hasGoodDevelopedCoping else 'Плохо')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Алекситимия: ')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph('Предполагается' if card.hasAlexithymiaPresumed else 'Не предполагается')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.alexithymiaNote:
        doc.add_paragraph(card.alexithymiaNote)
        c += 1
        doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Нарушения сна:')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.isHardToFallAsleep:
        doc.add_paragraph('\t- Нарушение засыпания')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasFrequentWaking:
        doc.add_paragraph('\t- Частое пробуждение')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.isAwakeEarly:
        doc.add_paragraph('\t- Просыпается рано ')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
        doc.paragraphs[c].add_run(card.earlyAwakeningNote if card.earlyAwakeningNote else '')
        doc.paragraphs[c].runs[1].font.size = Pt(12)
    if card.hasSleepDeprivation:
        doc.add_paragraph('\t- Чувство после сна, что не выспался')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasIncreasedDrowsiness:
        doc.add_paragraph('\t- Повышение сонливости: много спит и ночью и днем')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Справляется :')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.isCopeWithSleepDisturbancesActual.value)
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)


    # Раздел “Аппетит”
    doc.add_heading('Аппетит:', 2)
    c += 1
    doc.paragraphs[c].runs[0].bold = True
    doc.paragraphs[c].style = "Heading 2"
    doc.paragraphs[c].runs[0].font.size = Pt(14)

    doc.add_paragraph(card.appetiteDisordersActual.value)
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Пациент регулирует нарушения:')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.appetiteRegulationActionNote if card.appetiteRegulationActionNote else 'Нет.')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)


    # Раздел “Когнитивные функции”
    doc.add_heading('Когнитивные функции: ', 2)
    c += 1
    doc.paragraphs[c].runs[0].bold = True
    doc.paragraphs[c].style = "Heading 2"
    doc.paragraphs[c].runs[0].font.size = Pt(14)

    doc.add_paragraph('Понимание информации у пациента:')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.cognUnderstandingOfInfoNote if card.cognUnderstandingOfInfoNote else 'Нет.')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Запоминание информации:')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.cognMemoizationOfInfoNote if card.cognMemoizationOfInfoNote else 'Нет.')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Речь:')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.cognSpeechNote if card.cognSpeechNote else 'Нет.')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Праксис:')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.cognPraxisNote if card.cognPraxisNote else 'Нет.')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Социальный интеллект: ')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.cognSocIntelligenceNote if card.cognSocIntelligenceNote else 'Нет.')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('В беседе использовалась терапия:')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.useCognBehavTherapy:
        doc.add_paragraph('\t- Когнитивно-поведенческая')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.useExistentialTherapy:
        doc.add_paragraph('\t- Экзистенцильная')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.useWorkImagesTherapy:
        doc.add_paragraph('\t- Работа с внутренними образами')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.useDecisionTherapy:
        doc.add_paragraph('\t- Сфокусированная на решении')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.useAcceptanceTherapy:
        doc.add_paragraph('\t- Сфокусированная на принятии')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.useArtTherapy:
        doc.add_paragraph('\t- Арт-терапия')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.useBodilyOrientedTherapy:
        doc.add_paragraph('\t- Телесно-ориентированная')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.useOtherTherapy:
        doc.add_paragraph('\t- Другая. ')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
        doc.paragraphs[c].add_run(card.useOtherTherapyNote if card.useOtherTherapyNote else ' ')
        doc.paragraphs[c].runs[1].font.size = Pt(12)

    
    # Раздел “Заключение”
    doc.add_heading('Заключение: ', 2)
    c += 1
    doc.paragraphs[c].runs[0].bold = True
    doc.paragraphs[c].style = "Heading 2"
    doc.paragraphs[c].runs[0].font.size = Pt(14)

    doc.add_paragraph('На первый план выходит: ')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.conclFirstPlanNote if card.conclFirstPlanNote else ' ')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('На фоне:')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph(card.conclOnBackgroundNote if card.conclOnBackgroundNote else ' ')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Внутренняя картина болезни: ')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    doc.add_paragraph('Сформирована' if card.hasInnerPicture else 'Не сформирована')
    c += 1
    doc.paragraphs[c].runs[0].font.size = Pt(12)

    doc.add_paragraph('Предполагается: ')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasAdaptationDisorder:
        doc.add_paragraph('\t- Расстройство адаптации')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasAnxietyDisorder:
        doc.add_paragraph('\t- Тревожное расстройство')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasDepressionClinical:
        doc.add_paragraph('\t- Депрессия клиническая')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasAcutePsyReact:
        doc.add_paragraph('\t- Острая психологическая реакция')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasPostDisorder:
        doc.add_paragraph('\t- Посттравматическое расстройство')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.hasConclOther:
        doc.add_paragraph('\t- Другая. ')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
        doc.paragraphs[c].add_run(card.conclOtherNote)
        doc.paragraphs[c].runs[1].font.size = Pt(12)

    doc.add_paragraph(card.conclTargetNote if card.conclTargetNote else ' ')
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    c+=1

    doc.add_paragraph('Рекомендовано: ')
    c += 1
    doc.paragraphs[c].runs[0].italic = True
    doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.recomNeuroStudy:
        doc.add_paragraph('\t- Пройти дополнительное нейропсихологическое исследование')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.recomPathoStudy:
        doc.add_paragraph('\t- Пройти дополнительное патопсихологическое исследование')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.recomContactSpecialist:
        doc.add_paragraph('\t- Обратиться к врачу-психотерапевту/психиатру')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
    if card.recomPassTraining:
        doc.add_paragraph('\t- Пройти тренинг: ')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
        doc.paragraphs[c].add_run(card.recomPassTrainingNote if card.recomPassTrainingNote else ' ')
        doc.paragraphs[c].runs[1].font.size = Pt(12)
    if card.recomHomework:
        doc.add_paragraph('\t- Домашнее задание: ')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
        doc.paragraphs[c].add_run(card.recomHomeworkNote if card.recomHomeworkNote else ' ')
        doc.paragraphs[c].runs[1].font.size = Pt(12)
    if card.recomAdditionalPsychometry:
        doc.add_paragraph('\t- Пройти дополнительную психометрию: ')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)
        doc.paragraphs[c].add_run(card.recomAdditionalPsychometryNote if card.recomAdditionalPsychometryNote else ' ')
        doc.paragraphs[c].runs[1].font.size = Pt(12)
    if card.recomPsychocorrection:
        doc.add_paragraph('\t- Психокоррекция по запросу')
        c+=1
        doc.paragraphs[c].runs[0].font.size = Pt(12)

    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f